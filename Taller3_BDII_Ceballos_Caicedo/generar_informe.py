#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador del informe Word - Taller 3 BD II UIS 2026
Requiere: pip install python-docx
Ejecutar: python3 generar_informe.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE       = Path(__file__).parent
SCRIPTS    = BASE / 'scripts'
EVIDENCIAS = BASE / 'evidencias'
OUT        = BASE / 'informe' / 'Taller3_BDII_Ceballos_Caicedo.docx'

def buscar_imagen(fig_name, n):
    """Busca el archivo de imagen con nombre exacto, case-insensitive en extension."""
    base_name = f'fig{n:02d}_{fig_name}'
    for ext in ('.png', '.PNG', '.jpg', '.JPG', '.jpeg', '.JPEG'):
        p = EVIDENCIAS / (base_name + ext)
        if p.exists():
            return p
    return None

def leer(fname):
    with open(SCRIPTS / fname, encoding='utf-8') as f:
        return f.read()

# ============================================================
#  QUERIES PARA EL INFORME
# ============================================================

Q_INS_A = """\
MATCH (jefe:Empleado {codigoEmpresarial: 'EMP-008'})
CREATE (e:Empleado {
  codigoEmpresarial: 'EMP-019',
  documento:         '1068901234',
  nombre:            'Roberto',
  apellidos:         'Caicedo Sanchez',
  edad:              28,
  cargo:             'Empleado'
})
MERGE (jefe)-[:JEFE_DE]->(e)
RETURN e.codigoEmpresarial AS codigo,
       e.nombre + ' ' + e.apellidos AS nombreCompleto,
       e.cargo AS cargo,
       jefe.codigoEmpresarial AS codigoJefe;"""

Q_INS_B = """\
MATCH (nuevo:Empleado {codigoEmpresarial: 'EMP-019'})
MATCH (amigo:Empleado {codigoEmpresarial: 'EMP-014'})
MERGE (nuevo)-[:AMIGO_DE]->(amigo)
RETURN nuevo.nombre + ' ' + nuevo.apellidos AS empleado,
       amigo.nombre + ' ' + amigo.apellidos AS nuevoAmigo;"""

Q_INS_C = """\
MATCH (e:Empleado {codigoEmpresarial: 'EMP-019'})
CREATE (p:Post {
  id:          'POST-019',
  fecha:       datetime('2026-05-13T10:00:00'),
  titulo:      'Mi primer dia en la red corporativa',
  descripcion: 'Hola a todos. Me uno a la red social de la empresa con entusiasmo.'
})
MERGE (e)-[:PUBLICO]->(p)
RETURN p.id AS idPost, p.titulo AS titulo, e.nombre AS autor;"""

Q_INS_D = """\
MATCH (post:Post   {id:               'POST-019'})
MATCH (e1:Empleado {codigoEmpresarial: 'EMP-005'})
MATCH (e2:Empleado {codigoEmpresarial: 'EMP-011'})
CREATE (c1:Comentario {
  id:          'COM-037',
  fecha:       datetime('2026-05-13T10:30:00'),
  descripcion: 'Bienvenido Roberto. Un gusto tenerte en el equipo.'
})
CREATE (c2:Comentario {
  id:          'COM-038',
  fecha:       datetime('2026-05-13T11:00:00'),
  descripcion: 'Bienvenido. Cualquier duda nos cuentas sin problema.'
})
MERGE (e1)-[:COMENTO]->(c1)
MERGE (c1)-[:EN_POST]->(post)
MERGE (e2)-[:COMENTO]->(c2)
MERGE (c2)-[:EN_POST]->(post)
RETURN c1.id AS comentario1, e1.nombre AS autor1,
       c2.id AS comentario2, e2.nombre AS autor2;"""

Q_UPD_A = """\
MATCH (e:Empleado {codigoEmpresarial: 'EMP-016'})
SET e.edad  = 27,
    e.cargo = 'Administrador'
RETURN e.codigoEmpresarial AS codigo,
       e.nombre + ' ' + e.apellidos AS nombreCompleto,
       e.cargo AS nuevoCargo,
       e.edad  AS nuevaEdad;"""

Q_UPD_B = """\
MATCH (p:Post {id: 'POST-004'})
SET p.titulo      = 'Reunion de seguimiento - jueves 3 p.m.',
    p.descripcion = 'Este jueves a las 3 p.m. tenemos reunion de seguimiento. La sala de juntas quedo reservada. Confirmar asistencia antes del mediodia.'
RETURN p.id AS idPost, p.titulo AS nuevoTitulo;"""

Q_UPD_C = """\
MATCH (c:Comentario {id: 'COM-001'})
SET c.descripcion = 'Excelentes resultados Carlos Andres. El equipo trabajo con mucho compromiso. Estos numeros son reflejo del liderazgo y dedicacion de todos.'
RETURN c.id AS idComentario, c.descripcion AS nuevaDescripcion;"""

Q_DEL_A = """\
// Paso 1: eliminar comentarios publicados por el empleado
MATCH (e:Empleado {codigoEmpresarial: 'EMP-017'})-[:COMENTO]->(c:Comentario)
DETACH DELETE c;

// Paso 2: eliminar posts del empleado y sus comentarios asociados
MATCH (e:Empleado {codigoEmpresarial: 'EMP-017'})-[:PUBLICO]->(p:Post)
OPTIONAL MATCH (c:Comentario)-[:EN_POST]->(p)
DETACH DELETE c, p;

// Paso 3: eliminar reportes generados por el empleado
MATCH (e:Empleado {codigoEmpresarial: 'EMP-017'})-[:GENERA]->(r:Reporte)
DETACH DELETE r;

// Paso 4: eliminar el nodo Empleado y relaciones residuales
MATCH (e:Empleado {codigoEmpresarial: 'EMP-017'})
DETACH DELETE e;"""

Q_DEL_B = """\
// DETACH DELETE elimina el Reporte y sus relaciones GENERA y CONTRA.
// Los empleados vinculados (reportante y reportado) no se ven afectados.
MATCH (r:Reporte {id: 'REP-005'})
DETACH DELETE r;"""

Q_DEL_C = """\
// Paso 1: eliminar comentarios del post para no dejarlos huerfanos
MATCH (c:Comentario)-[:EN_POST]->(p:Post {id: 'POST-018'})
DETACH DELETE c;

// Paso 2: eliminar el post y sus relaciones restantes (DIO_LIKE, PUBLICO)
MATCH (p:Post {id: 'POST-018'})
DETACH DELETE p;"""

Q_RED_A = """\
MATCH (e:Empleado {cargo: $cargo})
RETURN e.codigoEmpresarial AS codigo,
       e.nombre + ' ' + e.apellidos AS nombreCompleto,
       e.edad AS edad,
       e.cargo AS cargo
ORDER BY nombreCompleto;
// Parametro sugerido: :param cargo => 'Administrador'"""

Q_RED_B = """\
MATCH (jefe:Empleado {codigoEmpresarial: $codigoEmpresarial})-[:JEFE_DE*1..]->(sub:Empleado)
RETURN DISTINCT sub.codigoEmpresarial AS codigo,
       sub.nombre + ' ' + sub.apellidos AS nombreCompleto,
       sub.cargo AS cargo
ORDER BY sub.cargo, nombreCompleto;
// Parametro sugerido: :param codigoEmpresarial => 'EMP-001'"""

Q_RED_C = """\
MATCH (e:Empleado {codigoEmpresarial: $codigoEmpresarial})-[:AMIGO_DE]-(amigo:Empleado)
RETURN amigo.codigoEmpresarial AS codigo,
       amigo.nombre + ' ' + amigo.apellidos AS nombreCompleto,
       amigo.cargo AS cargo
ORDER BY nombreCompleto;
// Parametro sugerido: :param codigoEmpresarial => 'EMP-005'"""

Q_RED_D = """\
MATCH (e:Empleado {codigoEmpresarial: $codigoEmpresarial})-[:AMIGO_DE]-(amigo:Empleado)
WITH e, collect(amigo) AS amigosDirectos
UNWIND amigosDirectos AS amigo
MATCH (amigo)-[:AMIGO_DE]-(amigodeamigo:Empleado)
WHERE amigodeamigo <> e
  AND NOT amigodeamigo IN amigosDirectos
RETURN DISTINCT amigodeamigo.codigoEmpresarial AS codigo,
       amigodeamigo.nombre + ' ' + amigodeamigo.apellidos AS nombreCompleto,
       amigodeamigo.cargo AS cargo
ORDER BY nombreCompleto;
// Parametro sugerido: :param codigoEmpresarial => 'EMP-005'"""

Q_RED_E = """\
MATCH (e:Empleado)-[:AMIGO_DE]-(amigo:Empleado)
WITH e, count(amigo) AS cantidadAmigos
ORDER BY cantidadAmigos DESC
LIMIT 1
RETURN e.codigoEmpresarial AS codigo,
       e.nombre + ' ' + e.apellidos AS nombreCompleto,
       e.cargo AS cargo,
       cantidadAmigos;"""

Q_POSTS_A = """\
MATCH (e:Empleado {codigoEmpresarial: $codigoEmpresarial})-[:PUBLICO]->(p:Post)
RETURN p.id AS idPost, p.fecha AS fecha,
       p.titulo AS titulo, p.descripcion AS descripcion
ORDER BY p.fecha DESC;
// Parametro sugerido: :param codigoEmpresarial => 'EMP-001'"""

Q_POSTS_B = """\
MATCH (c:Comentario)-[:EN_POST]->(p:Post {id: $idPost})
MATCH (autor:Empleado)-[:COMENTO]->(c)
RETURN c.id AS idComentario, c.fecha AS fecha,
       autor.nombre + ' ' + autor.apellidos AS autor,
       c.descripcion AS comentario
ORDER BY c.fecha ASC;
// Parametro sugerido: :param idPost => 'POST-001'"""

Q_POSTS_C = """\
MATCH (e:Empleado)-[:DIO_LIKE]->(p:Post {id: $idPost})
RETURN e.codigoEmpresarial AS codigo,
       e.nombre + ' ' + e.apellidos AS nombreCompleto,
       e.cargo AS cargo
ORDER BY nombreCompleto;
// Parametro sugerido: :param idPost => 'POST-001'"""

Q_POSTS_D = """\
MATCH (e:Empleado)-[:PUBLICO]->(p:Post)
OPTIONAL MATCH (:Empleado)-[like:DIO_LIKE]->(p)
WITH p, e, count(like) AS cantidadLikes
ORDER BY cantidadLikes DESC
LIMIT 3
RETURN p.id AS idPost, p.titulo AS titulo,
       e.nombre + ' ' + e.apellidos AS autor,
       cantidadLikes;"""

Q_ACOSO_A = """\
MATCH (reportado:Empleado {cargo: $cargo})<-[:CONTRA]-(r:Reporte)<-[:GENERA]-(reportante:Empleado)
RETURN reportado.codigoEmpresarial AS codigoReportado,
       reportado.nombre + ' ' + reportado.apellidos AS nombreReportado,
       reportado.cargo AS cargo,
       r.id AS idReporte,
       r.naturaleza AS naturaleza,
       r.fecha AS fechaReporte,
       reportante.nombre + ' ' + reportante.apellidos AS nombreReportante
ORDER BY r.fecha DESC;
// Parametro sugerido: :param cargo => 'Gerente'"""

Q_ACOSO_B = """\
MATCH (e:Empleado {codigoEmpresarial: $codigoEmpresarial})-[:GENERA]->(r:Reporte)
RETURN e.codigoEmpresarial AS codigo,
       e.nombre + ' ' + e.apellidos AS nombreEmpleado,
       count(r) AS cantidadReportesGenerados,
       collect(r.naturaleza) AS naturalezas;
// Parametro sugerido: :param codigoEmpresarial => 'EMP-011'"""

Q1 = """\
MATCH (e:Empleado {codigoEmpresarial: $codigoEmpresarial})-[:PUBLICO]->(p:Post)
RETURN p.id AS idPost, p.fecha AS fecha,
       p.titulo AS titulo, p.descripcion AS descripcion
ORDER BY p.fecha DESC;
// :param codigoEmpresarial => 'EMP-001'"""

Q2 = """\
MATCH (e:Empleado)-[:PUBLICO]->(p:Post)
RETURN p.id AS idPost, p.fecha AS fecha,
       p.titulo AS titulo,
       e.nombre + ' ' + e.apellidos AS autor
ORDER BY p.fecha DESC
LIMIT 3;"""

Q3 = """\
MATCH (e:Empleado)-[:PUBLICO]->(p:Post)
OPTIONAL MATCH (:Empleado)-[like:DIO_LIKE]->(p)
WITH p, e, count(like) AS cantidadLikes
ORDER BY cantidadLikes DESC
LIMIT 3
RETURN p.id AS idPost, p.titulo AS titulo,
       e.nombre + ' ' + e.apellidos AS autor,
       cantidadLikes;"""

Q4 = """\
MATCH (jefe:Empleado {codigoEmpresarial: $codigoEmpresarial})-[:JEFE_DE]->(sub:Empleado)
RETURN sub.codigoEmpresarial AS codigo,
       sub.nombre + ' ' + sub.apellidos AS nombreCompleto,
       sub.cargo AS cargo
ORDER BY nombreCompleto;
// :param codigoEmpresarial => 'EMP-001'"""

Q5 = """\
MATCH (e:Empleado {codigoEmpresarial: $codigoEmpresarial})-[:DIO_LIKE]->(p:Post)
MATCH (autor:Empleado)-[:PUBLICO]->(p)
RETURN p.id AS idPost, p.titulo AS titulo, p.fecha AS fecha,
       autor.nombre + ' ' + autor.apellidos AS autorPost
ORDER BY p.fecha DESC;
// :param codigoEmpresarial => 'EMP-005'"""

Q6 = """\
MATCH (reportado:Empleado {cargo: $cargo})<-[:CONTRA]-(r:Reporte)<-[:GENERA]-(reportante:Empleado)
RETURN reportado.nombre + ' ' + reportado.apellidos AS nombreReportado,
       reportado.cargo AS cargo,
       r.id AS idReporte,
       r.naturaleza AS naturaleza,
       reportante.nombre + ' ' + reportante.apellidos AS nombreReportante
ORDER BY r.fecha DESC;
// :param cargo => 'Gerente'"""

Q7 = """\
// Seleccionar vista "Graph" en Neo4j Browser para ver el grafo de amistades.
MATCH (e:Empleado {codigoEmpresarial: $codigoEmpresarial})-[rel:AMIGO_DE]-(amigo:Empleado)
RETURN e, rel, amigo;
// :param codigoEmpresarial => 'EMP-005'"""

Q8 = """\
MATCH (e:Empleado)-[:AMIGO_DE]-(amigo:Empleado)
WITH e, count(amigo) AS cantidadAmigos
ORDER BY cantidadAmigos DESC
LIMIT 3
RETURN e.codigoEmpresarial AS codigo,
       e.nombre + ' ' + e.apellidos AS nombreCompleto,
       e.cargo AS cargo,
       cantidadAmigos;"""

Q9 = """\
MATCH (c:Comentario)-[:EN_POST]->(p:Post)
MATCH (autor:Empleado)-[:PUBLICO]->(p)
WITH p, autor, count(c) AS cantidadComentarios
ORDER BY cantidadComentarios DESC
RETURN p.id AS idPost, p.titulo AS titulo,
       autor.nombre + ' ' + autor.apellidos AS autor,
       cantidadComentarios;"""

# ============================================================
#  INICIALIZAR DOCUMENTO
# ============================================================

doc = Document()
sec = doc.sections[0]
for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(sec, attr, Cm(2.5))
sec.different_first_page_header_footer = True

# ============================================================
#  HELPERS
# ============================================================

def rfmt(run, name='Calibri', size=11, bold=False, italic=False):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic

def pfmt(p, sb=0, sa=6, ls=1.15):
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = ls

def body(text='', align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    pfmt(p)
    if text:
        r = p.add_run(text)
        rfmt(r)
    return p

def h1(text):
    p = doc.add_paragraph()
    pfmt(p, sb=14, sa=6)
    r = p.add_run(text)
    rfmt(r, size=14, bold=True)
    return p

def h2(text):
    p = doc.add_paragraph()
    pfmt(p, sb=10, sa=4)
    r = p.add_run(text)
    rfmt(r, size=12, bold=True)
    return p

def h3(text):
    p = doc.add_paragraph()
    pfmt(p, sb=8, sa=3)
    r = p.add_run(text)
    rfmt(r, size=11, bold=True)
    return p

def codigo(text):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.cell(0, 0)
    c.paragraphs[0].clear()
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
    doc.add_paragraph()
    return t

_fig = [0]

def evidencia(fig_name, desc, query, vista):
    _fig[0] += 1
    n   = _fig[0]
    fn  = f'fig{n:02d}_{fig_name}.png'
    img = buscar_imagen(fig_name, n)

    if img:
        # Insertar imagen real
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after  = Pt(0)
        run = p_img.add_run()
        run.add_picture(str(img), width=Cm(14))
    else:
        # Placeholder cuando no existe la imagen
        t = doc.add_table(rows=1, cols=1)
        t.style = 'Table Grid'
        c = t.cell(0, 0)
        c.paragraphs[0].clear()
        lineas = [
            f'[EVIDENCIA - {fn}]',
            f'Descripcion:  {desc}',
            f'Query:        {query}',
            f'Vista:        {vista}',
            f'Archivo:      evidencias/{fn}',
        ]
        for i, ln in enumerate(lineas):
            p = c.paragraphs[0] if i == 0 else c.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            r = p.add_run(ln)
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
            r.font.bold = (i == 0)

    # Caption siempre presente
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after  = Pt(10)
    r = cap.add_run(f'Figura {n}. {desc}')
    rfmt(r, size=10, italic=True)
    return n, fn

def tabla_datos(headers, filas):
    t = doc.add_table(rows=len(filas) + 1, cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h)
        rfmt(r, bold=True)
    for i, fila in enumerate(filas):
        for j, val in enumerate(fila):
            c = t.cell(i + 1, j)
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(str(val))
            rfmt(r)
    doc.add_paragraph()
    return t

# ============================================================
#  FOOTER
# ============================================================

def setup_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()

    pPr = p._p.get_or_add_pPr()
    tabs_el = OxmlElement('w:tabs')
    tab_el  = OxmlElement('w:tab')
    tab_el.set(qn('w:val'), 'right')
    tab_el.set(qn('w:pos'), '9072')
    tabs_el.append(tab_el)
    pPr.append(tabs_el)

    r1 = p.add_run('Bases de Datos II – Taller N° 3   |   UIS 2026')
    rfmt(r1, size=9)
    p.add_run('\t')
    r2 = p.add_run('Página ')
    rfmt(r2, size=9)

    r3 = p.add_run()
    rfmt(r3, size=9)
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    r3._r.append(fc1)
    ins = OxmlElement('w:instrText')
    ins.set(qn('xml:space'), 'preserve')
    ins.text = 'PAGE'
    r3._r.append(ins)
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'end')
    r3._r.append(fc2)

setup_footer(sec)

# ============================================================
#  PORTADA / ENCABEZADO
# ============================================================

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pfmt(p, sb=0, sa=4)
r = p.add_run('BASES DE DATOS II')
rfmt(r, size=18, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pfmt(p, sb=0, sa=4)
r = p.add_run('Universidad Industrial de Santander')
rfmt(r, size=13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pfmt(p, sb=0, sa=4)
r = p.add_run('TALLER N° 3')
rfmt(r, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pfmt(p, sb=0, sa=12)
r = p.add_run('Diseño e implementación de una base de datos NoSQL orientada a grafos en Neo4j')
rfmt(r, size=12)

tbl = doc.add_table(rows=4, cols=2)
tbl.style = 'Table Grid'
meta = [
    ('Estudiantes', 'Kaleth Ceballos\nRobert Caicedo'),
    ('Grupo',       'E1'),
    ('Docente',     'Jathinson Meneses Mendoza'),
    ('Fecha',       'Mayo de 2026'),
]
for i, (k, v) in enumerate(meta):
    for j, txt in enumerate((k, v)):
        c = tbl.cell(i, j)
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(txt)
        rfmt(r, bold=(j == 0))
doc.add_paragraph()

# ============================================================
#  ENUNCIADO DEL PROBLEMA
# ============================================================

h1('Enunciado del Problema')
body(
    'Una empresa colombiana con empleados distribuidos en todo el país requiere '
    'construir una red social corporativa. Los empleados poseen muros personales donde '
    'publican contenido; otros colaboradores pueden dar me gusta y comentar dichas '
    'publicaciones, y también pueden agregarse como amigos. La empresa cuenta con una '
    'cadena de mando jerárquica estricta: Gerente, Líder Regional, Líder Seccional, '
    'Administrador y Empleado, en ese orden descendente. La relación JEFE_DE solo puede '
    'existir entre cargos inmediatamente adyacentes en esa cadena.'
)
body(
    'El comité de ética requiere mapear los reportes de acoso laboral para detectar '
    'patrones donde personas en cargos altos abusen de su posición. La naturaleza del '
    'acoso puede clasificarse en: maltrato, persecución, discriminación, entorpecimiento, '
    'inequidad, desprotección o acoso sexual. Se solicita implementar este modelo usando '
    'Neo4j Community Edition 5.x con lenguaje Cypher.'
)

# ============================================================
#  ATRIBUTOS DE LAS ENTIDADES
# ============================================================

h1('Atributos de las entidades')
tabla_datos(
    ['Entidad', 'Atributos'],
    [
        ('Empleado',    'codigoEmpresarial (PK), documento (unico), nombre, apellidos, edad, cargo'),
        ('Post',        'id (PK), fecha, titulo, descripcion'),
        ('Comentario',  'id (PK), fecha, descripcion'),
        ('Reporte',     'id (PK), fecha, naturaleza'),
    ]
)
body(
    'La cantidad de me gusta de un Post se deriva contando relaciones :DIO_LIKE entrantes; '
    'no es un atributo almacenado en el nodo.'
)

# ============================================================
#  RELACIONES DEL MODELO
# ============================================================

h1('Relaciones del modelo')
tabla_datos(
    ['Tipo de relacion', 'Origen → Destino', 'Descripcion'],
    [
        (':JEFE_DE',   'Empleado → Empleado',   'Jerarquia de mando. Solo entre cargos adyacentes en la cadena.'),
        (':AMIGO_DE',  'Empleado → Empleado',   'Amistad corporativa. Arista unica; consultas sin direccion.'),
        (':PUBLICO',   'Empleado → Post',        'Autoria del post.'),
        (':DIO_LIKE',  'Empleado → Post',        'Me gusta dado por el empleado.'),
        (':COMENTO',   'Empleado → Comentario',  'Autoria del comentario.'),
        (':EN_POST',   'Comentario → Post',      'Comentario pertenece al post.'),
        (':GENERA',    'Empleado → Reporte',     'Empleado que interpone el reporte.'),
        (':CONTRA',    'Reporte → Empleado',     'Empleado que es reportado.'),
    ]
)

# ============================================================
#  1. DISENO DEL ESQUEMA DE GRAFO
# ============================================================

h1('1. Diseño del esquema de grafo')

h2('1.1 Modelo conceptual')
body(
    'El modelo de grafo representó cada concepto del dominio como un nodo etiquetado. '
    'Se optaron por nodos separados para Comentario y Reporte en lugar de propiedades '
    'de las relaciones, porque ambas entidades poseen atributos propios relevantes para '
    'consultas (fecha, descripción, naturaleza) y pueden participar en múltiples '
    'relaciones. Esta decisión facilita el traversal y la búsqueda eficiente.'
)
body(
    'La cadena de mando se modeló con la relación :JEFE_DE, restringida a cargos '
    'adyacentes. Esto permite recorrer la jerarquía a cualquier profundidad con patrones '
    'de longitud variable ([:JEFE_DE*1..]), algo que en un modelo relacional '
    'requeriría consultas recursivas o tablas de cierre transitivo.'
)
body(
    'La amistad se modeló como una sola arista dirigida :AMIGO_DE, siguiendo el patrón '
    'canónico de Neo4j para relaciones simétricas. Las consultas omiten la dirección '
    'usando el patrón (a)-[:AMIGO_DE]-(b), lo que evita duplicar aristas y mantiene '
    'el grafo limpio.'
)

h2('1.2 Diagrama del esquema')
evidencia(
    'esquema',
    'Esquema del grafo generado por Neo4j',
    'CALL db.schema.visualization()',
    'Grafo'
)

h2('1.3 Constraints e índices')
body('El siguiente script define los constraints de unicidad, NOT NULL e índices del esquema:')
codigo(leer('01_schema_constraints.cypher'))

h2('1.4 Justificación del modelo')
body(
    'Se crearon nodos separados para Comentario y Reporte porque tienen ciclo de vida '
    'independiente del Post o del Empleado y deben ser consultables de forma autónoma '
    '(ej. todos los comentarios de un post, todos los reportes con naturaleza discriminacion). '
    'La alternativa de representarlos como propiedades habría obligado a usar listas '
    'y perdido capacidad de traversal.'
)
body(
    'La relación :AMIGO_DE como arista única (sin duplicar en sentido inverso) sigue el '
    'principio de normalidad del grafo: el patrón sin dirección (a)-[:AMIGO_DE]-(b) cubre '
    'ambos sentidos en la consulta. Duplicar la arista generaría resultados dobles en '
    'conteos de amigos y consultas de amigos de amigos.'
)
body(
    'Los constraints NOT NULL garantizan integridad desde el motor; los índices sobre '
    'Empleado.cargo, Post.fecha y Reporte.fecha responden a los predicados de filtrado '
    'más frecuentes en las consultas del enunciado.'
)

# ============================================================
#  2. CARGA DE DATOS
# ============================================================

h1('2. Carga de datos')

h2('2.1 Estrategia de carga')
body(
    'Los datos se prepararon en 11 archivos CSV con encabezados, codificación UTF-8 sin '
    'BOM y fechas en formato ISO 8601 (YYYY-MM-DDTHH:MM:SS). Los archivos se copiaron '
    'a la carpeta import de la base de datos de Neo4j Desktop antes de ejecutar el '
    'script de carga.'
)
body(
    'El script 02_carga_csv.cypher utiliza LOAD CSV WITH HEADERS y MERGE para cada '
    'entidad, garantizando idempotencia: puede re-ejecutarse sin crear duplicados. '
    'La carga sigue el orden: primero los nodos (Empleado, Post, Comentario, Reporte) '
    'y luego las relaciones, respetando la dependencia entre ellos.'
)

h2('2.2 Volumen de datos cargados')
tabla_datos(
    ['Tipo', 'Label / Tipo de relacion', 'Cantidad'],
    [
        ('Nodo',      'Empleado',    '18'),
        ('Nodo',      'Post',        '18'),
        ('Nodo',      'Comentario',  '36'),
        ('Nodo',      'Reporte',     '5'),
        ('Relacion',  ':JEFE_DE',    '17'),
        ('Relacion',  ':AMIGO_DE',   '24'),
        ('Relacion',  ':PUBLICO',    '18'),
        ('Relacion',  ':DIO_LIKE',   '61'),
        ('Relacion',  ':COMENTO',    '36'),
        ('Relacion',  ':EN_POST',    '36'),
        ('Relacion',  ':GENERA',     '5'),
        ('Relacion',  ':CONTRA',     '5'),
        ('TOTAL NODOS',      '',     '77'),
        ('TOTAL RELACIONES', '',     '207'),
    ]
)

h2('2.3 Script de carga')
codigo(leer('02_carga_csv.cypher'))

h2('2.4 Evidencia - Grafo completo')
evidencia(
    'grafo_completo',
    'Grafo completo despues de la carga de datos',
    'MATCH (n) RETURN n LIMIT 300',
    'Grafo'
)

# ============================================================
#  3. OPERACIONES DE MANIPULACION
# ============================================================

h1('3. Operaciones de manipulación')

h2('3.1 Inserción')

h3('a) Insertar un nuevo usuario')
body(
    'Se crea el empleado Roberto Caicedo Sánchez (EMP-019) con cargo Empleado, '
    'y se vincula como subordinado del Administrador Miguel Angel Acevedo Torres (EMP-008).'
)
codigo(Q_INS_A)
evidencia('insercion_a', 'Resultado de insercion del empleado EMP-019',
          "Ejecutar bloque a) de 03_queries_insercion.cypher", 'Tabla')

h3('b) Insertar una nueva relacion de amistad')
body('Se crea una relación :AMIGO_DE entre EMP-019 y EMP-014 (Edwin Fabian Zapata Giraldo).')
codigo(Q_INS_B)
evidencia('insercion_b', 'Resultado de insercion de amistad EMP-019 con EMP-014',
          "Ejecutar bloque b) de 03_queries_insercion.cypher", 'Tabla')

h3('c) Insertar un nuevo post')
body('EMP-019 publica su primer post (POST-019) en la red social corporativa.')
codigo(Q_INS_C)
evidencia('insercion_c', 'Resultado de insercion del post POST-019',
          "Ejecutar bloque c) de 03_queries_insercion.cypher", 'Tabla')

h3('d) Insertar comentarios de otros empleados')
body(
    'EMP-005 (Andres Felipe Mena Palacios) y EMP-011 (Natalia Andrea Cuellar Fonseca) '
    'comentan el post POST-019.'
)
codigo(Q_INS_D)
evidencia('insercion_d', 'Resultado de insercion de comentarios COM-037 y COM-038',
          "Ejecutar bloque d) de 03_queries_insercion.cypher", 'Tabla')

h2('3.2 Actualización')

h3('a) Actualizar informacion de un empleado')
body(
    'Se actualiza la edad y el cargo de EMP-016 (Ferney Andres Chavez Beltran), '
    'que asciende a Administrador.'
)
codigo(Q_UPD_A)
evidencia('actualizacion_a', 'Resultado de actualizacion del empleado EMP-016',
          "Ejecutar bloque a) de 04_queries_update.cypher", 'Tabla')

h3('b) Actualizar el contenido de un POST')
body('Se corrige el título y la descripción de POST-004.')
codigo(Q_UPD_B)
evidencia('actualizacion_b', 'Resultado de actualizacion del post POST-004',
          "Ejecutar bloque b) de 04_queries_update.cypher", 'Tabla')

h3('c) Actualizar un comentario')
body('Se amplia la descripción del comentario COM-001.')
codigo(Q_UPD_C)
evidencia('actualizacion_c', 'Resultado de actualizacion del comentario COM-001',
          "Ejecutar bloque c) de 04_queries_update.cypher", 'Tabla')

h2('3.3 Eliminación')
body(
    'NOTA: Ejecutar el script 00_reset_db.cypher y luego 02_carga_csv.cypher para '
    'restaurar los datos originales despues de correr las eliminaciones.'
)

h3('a) Eliminar un empleado')
body(
    'Se elimina EMP-017 (Yolanda Maria Villanueva Arias) siguiendo el orden de borrado '
    'responsable: primero sus comentarios, luego sus posts con comentarios asociados, '
    'luego sus reportes generados, y finalmente el nodo Empleado con DETACH DELETE.'
)
codigo(Q_DEL_A)
evidencia('eliminacion_a',
          'Verificacion de eliminacion de EMP-017',
          "MATCH (e:Empleado {codigoEmpresarial: 'EMP-017'}) RETURN e",
          'Tabla')

h3('b) Eliminar un reporte de acoso laboral')
body(
    'Se elimina REP-005 (naturaleza: inequidad). DETACH DELETE remueve el nodo Reporte '
    'y sus relaciones :GENERA y :CONTRA. Los empleados vinculados no se afectan.'
)
codigo(Q_DEL_B)
evidencia('eliminacion_b',
          'Verificacion de eliminacion del reporte REP-005',
          "MATCH (r:Reporte {id: 'REP-005'}) RETURN r",
          'Tabla')

h3('c) Eliminar un POST')
body(
    'Se elimina POST-018. Primero se borran los Comentario asociados via :EN_POST para '
    'no dejarlos como nodos huérfanos; luego se borra el Post con DETACH DELETE.'
)
codigo(Q_DEL_C)
evidencia('eliminacion_c',
          'Verificacion de eliminacion del post POST-018',
          "MATCH (p:Post {id: 'POST-018'}) RETURN p",
          'Tabla')

# ============================================================
#  4. CONSULTAS DE BUSQUEDA
# ============================================================

h1('4. Consultas de búsqueda')

h2('4.1 Red social y jerarquía')

h3('a) Empleados que desempeñan un rol determinado')
body('Filtra todos los empleados cuyo cargo coincide con el parámetro dado.')
codigo(Q_RED_A)
evidencia('red_por_cargo',
          'Empleados con cargo Administrador',
          "Ejecutar query a) de 06_queries_search_red.cypher con cargo='Administrador'",
          'Tabla')

h3('b) Empleados a cargo de un empleado dado (todos los niveles)')
body('Recorre :JEFE_DE con longitud variable para obtener toda la cadena de subordinados.')
codigo(Q_RED_B)
evidencia('red_subordinados',
          'Todos los subordinados del Gerente EMP-001',
          "Ejecutar query b) de 06_queries_search_red.cypher con codigoEmpresarial='EMP-001'",
          'Tabla')

h3('c) Amigos de un empleado')
body('El patrón sin dirección cubre la arista en ambos sentidos.')
codigo(Q_RED_C)
evidencia('red_amigos',
          'Amigos de EMP-005 (Andres Felipe Mena Palacios)',
          "Ejecutar query c) de 06_queries_search_red.cypher con codigoEmpresarial='EMP-005'",
          'Tabla')

h3('d) Amigos de los amigos de un empleado')
body('Excluye al empleado consultado y a sus amigos directos del resultado.')
codigo(Q_RED_D)
evidencia('red_amigos_amigos',
          'Amigos de los amigos de EMP-005, excluyendo directos',
          "Ejecutar query d) de 06_queries_search_red.cypher con codigoEmpresarial='EMP-005'",
          'Tabla')

h3('e) Empleado con más amigos')
codigo(Q_RED_E)
evidencia('red_mas_amigos',
          'Empleado con mayor cantidad de amigos en el grafo',
          "Ejecutar query e) de 06_queries_search_red.cypher",
          'Tabla')

h2('4.2 Posts y comentarios')

h3('a) POSTs de un empleado determinado')
codigo(Q_POSTS_A)
evidencia('posts_por_empleado',
          'Posts publicados por EMP-001 (Carlos Andres Morales Gutierrez)',
          "Ejecutar query a) de 07_queries_search_posts.cypher con codigoEmpresarial='EMP-001'",
          'Tabla')

h3('b) Comentarios de un POST')
body('Incluye el autor de cada comentario y se ordena cronológicamente.')
codigo(Q_POSTS_B)
evidencia('comentarios_post',
          'Comentarios del post POST-001 con sus autores',
          "Ejecutar query b) de 07_queries_search_posts.cypher con idPost='POST-001'",
          'Tabla')

h3('c) Empleados que dieron me gusta a un POST')
codigo(Q_POSTS_C)
evidencia('likes_post',
          'Empleados que dieron me gusta al post POST-001',
          "Ejecutar query c) de 07_queries_search_posts.cypher con idPost='POST-001'",
          'Tabla')

h3('d) Los 3 POSTs con más me gusta')
body(
    'Se usa OPTIONAL MATCH para incluir posts con cero likes y '
    'se ordena antes del LIMIT para garantizar el top correcto.'
)
codigo(Q_POSTS_D)
evidencia('top3_likes',
          'Los 3 posts con mayor cantidad de me gusta',
          "Ejecutar query d) de 07_queries_search_posts.cypher",
          'Tabla')

h2('4.3 Reportes de acoso')

h3('a) Empleados con reportes de acoso laboral dado un cargo')
body(
    'Retorna el empleado reportado, su cargo, el tipo de acoso, '
    'la fecha del reporte y quién lo interpuso.'
)
codigo(Q_ACOSO_A)
evidencia('acoso_por_cargo',
          'Reportes de acoso contra empleados con cargo Gerente',
          "Ejecutar query a) de 08_queries_search_acoso.cypher con cargo='Gerente'",
          'Tabla')

h3('b) Número de reportes de acoso realizados por un empleado')
codigo(Q_ACOSO_B)
evidencia('reportes_por_empleado',
          'Reportes generados por EMP-011 (Natalia Andrea Cuellar Fonseca)',
          "Ejecutar query b) de 08_queries_search_acoso.cypher con codigoEmpresarial='EMP-011'",
          'Tabla')

# ============================================================
#  5. CONSULTAS PRINCIPALES DEL ANALISIS
# ============================================================

h1('5. Consultas principales del análisis')
body(
    'A continuación se presentan las nueve consultas principales del enunciado '
    '(Diapositiva 7). El archivo 09_queries_principales.cypher las contiene '
    'de forma autocontenida con parámetros sugeridos.'
)

h2('Q1. Listado de POSTs de un empleado determinado')
codigo(Q1)
evidencia('q1_posts_empleado',
          'Posts publicados por el Gerente EMP-001',
          "Q1 de 09_queries_principales.cypher con codigoEmpresarial='EMP-001'",
          'Tabla')

h2('Q2. Los 3 POSTs más recientes publicados')
codigo(Q2)
evidencia('q2_posts_recientes',
          'Los 3 posts mas recientes de la red social',
          "Q2 de 09_queries_principales.cypher",
          'Tabla')

h2('Q3. Los 3 POSTs con más me gusta')
body(
    'POST-001 (12 likes), POST-002 (10 likes) y POST-003 (9 likes) deben encabezar '
    'el resultado según la distribución de datos cargados.'
)
codigo(Q3)
evidencia('q3_posts_likes',
          'Top 3 posts con mas me gusta - POST-001 (12), POST-002 (10), POST-003 (9)',
          "Q3 de 09_queries_principales.cypher",
          'Tabla')

h2('Q4. Empleados que tienen como jefe a un empleado determinado')
body('Retorna los subordinados directos (un nivel) del empleado indicado.')
codigo(Q4)
evidencia('q4_subordinados_directos',
          'Subordinados directos del Gerente EMP-001',
          "Q4 de 09_queries_principales.cypher con codigoEmpresarial='EMP-001'",
          'Tabla')

h2('Q5. Listado de POSTs a los que un empleado le ha dado me gusta')
codigo(Q5)
evidencia('q5_posts_likeados',
          'Posts que le gustaron a EMP-005 (Andres Felipe Mena Palacios)',
          "Q5 de 09_queries_principales.cypher con codigoEmpresarial='EMP-005'",
          'Tabla')

h2('Q6. Empleados con reportes de acoso en un cargo específico')
body(
    'Se verifican los dos reportes contra cargos altos: REP-001 contra el Gerente '
    '(persecución) y REP-002 contra el Líder Regional (discriminación).'
)
codigo(Q6)
evidencia('q6_acoso_cargo',
          'Empleados reportados por acoso con cargo Gerente',
          "Q6 de 09_queries_principales.cypher con cargo='Gerente'",
          'Tabla')

h2('Q7. Lista de amigos de un empleado mostrando relaciones en el grafo')
body(
    'Seleccionar la vista "Graph" en Neo4j Browser para visualizar los nodos y '
    'aristas :AMIGO_DE. La vista de tabla solo muestra atributos, no el grafo.'
)
codigo(Q7)
evidencia('q7_amigos_grafo',
          'Grafo de amistades de EMP-005 (6 amigos) - vista Graph en Neo4j Browser',
          "Q7 de 09_queries_principales.cypher con codigoEmpresarial='EMP-005'",
          'Grafo')

h2('Q8. Los 3 empleados con más amigos')
body(
    'Resultado esperado: EMP-005 (6 amigos), EMP-011 (6 amigos), EMP-014 (5 amigos).'
)
codigo(Q8)
evidencia('q8_top3_amigos',
          'Top 3 empleados con mas amigos - EMP-005 y EMP-011 (6) y EMP-014 (5)',
          "Q8 de 09_queries_principales.cypher",
          'Tabla')

h2('Q9. POSTs con más comentarios')
body('Todos los posts tienen exactamente 2 comentarios según los datos cargados.')
codigo(Q9)
evidencia('q9_posts_comentarios',
          'Posts con mayor cantidad de comentarios',
          "Q9 de 09_queries_principales.cypher",
          'Tabla')

# ============================================================
#  6. CONCLUSIONES
# ============================================================

h1('6. Conclusiones')
body(
    'El modelo orientado a grafos resultó idóneo para representar la red social '
    'corporativa. La naturaleza del dominio —relaciones de amistad, cadenas de mando, '
    'interacciones sobre publicaciones— es inherentemente un grafo, y Neo4j expresó '
    'estas conexiones de forma directa sin uniones artificiales. Consultas que en SQL '
    'requerirían múltiples JOINs recursivos (como obtener todos los subordinados de '
    'cualquier nivel) se resolvieron en Cypher con un único patrón de longitud variable.'
)
body(
    'La cadena de mando jerárquica demandó especial cuidado en la carga de datos: '
    'validar que cada relación :JEFE_DE conectara cargos adyacentes fue una '
    'responsabilidad de la lógica de preparación de los CSV, no del motor. '
    'En un entorno de producción, esta validación deberia implementarse en la capa '
    'de aplicación o mediante procedimientos almacenados, ya que Neo4j Community '
    'Edition no soporta triggers nativos.'
)
body(
    'La detección de patrones de acoso laboral mediante grafos ofrece ventajas '
    'significativas frente a una base de datos relacional. Con Cypher es posible '
    'extender las consultas para identificar si el reportado tiene autoridad jerárquica '
    'sobre el reportante, cruzando los patrones :CONTRA y :JEFE_DE en una sola '
    'consulta. Este tipo de análisis multi-hop es complejo en SQL pero natural '
    'en el modelo de grafo.'
)
body(
    'En cuanto a escalabilidad, Neo4j maneja bien grafos con millones de nodos, '
    'y los índices definidos sobre Empleado.cargo, Post.fecha y Reporte.fecha '
    'garantizan acceso directo a los predicados de filtrado más frecuentes. '
    'Sin embargo, para queries analíticos masivos (como conteos globales o '
    'agrupaciones sobre todo el grafo) los motores de grafos en memoria o '
    'soluciones OLAP serían complementos naturales.'
)
body(
    'El proyecto demostró la madurez de Neo4j 5.x para modelar dominios sociales: '
    'constraints NOT NULL, unicidad, índices y la sintaxis declarativa de Cypher '
    'permiten construir un esquema robusto con validación a nivel de motor, lo que '
    'reduce la superficie de errores de integridad en comparación con otras bases '
    'de datos NoSQL como MongoDB o Cassandra.'
)

# ============================================================
#  GUARDAR
# ============================================================

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f'Informe generado en: {OUT}')
print(f'Total de figuras a capturar: {_fig[0]}')
